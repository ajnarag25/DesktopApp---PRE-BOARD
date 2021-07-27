from PIC1 import *
from PIC2 import *
from PIC3 import *
from PIC4 import *
from PIC5 import *
from PIC6 import *
from MAIN import *
from STUDENT import *
from FIRST import *
from LOGIN import *
from ADMINLOGIN import *
from SIGNUP import *
from EXAMINEES import *
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import *
import sys
import datetime
import mysql.connector
import os
from docx import Document
import xlsxwriter


class First(First, QMainWindow):
    def __init__(self):
        super(First, self).__init__()

        self.setupUi(self)
        self.cone = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")
        self.Student.clicked.connect(self.Estudyante)
        self.Admin.clicked.connect(self.ADMIN)

    def Estudyante(self):
        una.hide()
        sign.hide()
        system.hide()
        log.show()

    def ADMIN(self):
        una.hide()
        sign.hide()
        system.hide()
        adm.show()


class Admin(Admin, QMainWindow):
    def __init__(self):
        super(Admin, self).__init__()

        self.setupUi(self)

        self.ImNotAnAdmin.clicked.connect(self.Notadmin)
        self.CONFIRM.clicked.connect(self.success)

    def Notadmin(self):
        una.show()
        adm.hide()
    def success(self):
        Userr = self.Username.text()
        Passw = self.Password.text()
        if Userr == 'Admin' and Passw == 'admin':
            system.show()
            self.Username.clear()
            self.Password.clear()
            adm.hide()
        elif Userr == '' and Passw == '':
            QMessageBox.warning(self, 'Dont leave it blank!', 'Invalid Username/Password')
        else:
            QMessageBox.warning(self,'Admin Account Could not verify','Invalid Username/Password')


class Login(Login, QMainWindow):
    def __init__(self):
        super(Login, self).__init__()

        self.setupUi(self)

        self.CONFIRM.clicked.connect(self.Confirm)
        self.SIGNUP.clicked.connect(self.Signup)
        self.RETURN.clicked.connect(self.Return)

        self.cone = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")


    def Confirm(self):
        User = self.Username.text()
        Pass = self.Password.text()
        cursor1=self.cone.cursor()

        query = "select * from naragdata where username=%s and passwd=%s;"
        data = (User, Pass)
        print(data)
        cursor1.execute(query, data)
        a = cursor1.fetchall()
        print(a)
        length = len(a)
        if length > 0:
            estudyante.show()
            log.hide()
            self.Username.clear()
            self.Password.clear()
            query1 = "select stat from naragdata where username=%s and passwd=%s;"
            data = (User, Pass)
            cursor1.execute(query1, data)
            a = cursor1.fetchall()
            print(a)
            estudyante.THERESULT.setText(str(a))
        elif User == '' and Pass == '':
            QMessageBox.warning(self, "Dont leave it blank!", "Invalid Username/Password")
        else:
            QMessageBox.warning(self, "User is not in the database!", "Invalid Username/Password")
            self.Username.clear()
            self.Password.clear()

        self.cone.commit()



    def Signup(self):
        una.hide()
        sign.show()
        system.hide()
        log.hide()

    def Return(self):
        log.hide()
        una.show()



class Signup(Signup, QMainWindow):
    def __init__(self):
        super(Signup, self).__init__()
        self.setupUi(self)
        self.conn = mysql.connector.connect(host="localhost", user="root", password="narag", database="preboard")
        self.ConfirmButton_2.clicked.connect(self.Back)
        self.ConfirmButton.clicked.connect(self.signupka)


    def signupka(self):
        lastt = self.Lastname.text()
        firstt = self.Firstname.text()
        midd = self.Middlename.text()
        usern = self.Username.text()
        passw = self.Password.text()
        conf = self.ConfirmPassword.text()
        if self.firstshift.isChecked()==True:
            radio="Firstshift"
        elif self.secondshift.isChecked()==True:
            radio="Secondshift"

        if passw != conf:
            QMessageBox.warning(self,'Invalid','Please Check Your Password')
        elif lastt == '' or firstt == '' or midd == '' or usern == '' or passw == '':
            QMessageBox.warning(self, 'Invalid', 'Please Dont Leave It Blank')
        elif self.firstshift.isChecked() == False and self.secondshift.isChecked() == False:
            QMessageBox.warning(self, 'Invalid', 'Please Select your shift')
        else:
            QMessageBox.about(self, "YOUR DATA HAS BEEN SAVED",
                              "PLEASE CLICK THE 'BACK TO LOGIN BUTTON' TO LOGIN YOUR ACCOUNT")
            cursor = self.conn.cursor()
            query = "insert into naragdata(lastname, firstname, middlename,shift, username, passwd) values(%s, %s, %s, %s, %s, %s);"
            data = (lastt, firstt, midd,radio, usern, passw)
            cursor.execute(query, data)
            self.conn.commit()
            self.Lastname.clear()
            self.Firstname.clear()
            self.Middlename.clear()
            self.Username.clear()
            self.Password.clear()
            self.firstshift.clearMask()
            self.ConfirmPassword.clear()

    def Back(self):
        una.hide()
        sign.hide()
        system.hide()
        log.show()


class Student(Student, QMainWindow):
    def __init__(self):
        super(Student, self).__init__()
        self.setupUi(self)


        self.cone = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")
        self.EXAMframe.hide()
        self.HOMEframe.show()
        self.StudentResults.hide()
        current = datetime.datetime.now()
        self.DATE.setText(str(current))


        self.EXAM.clicked.connect(self.ExamFrame)
        self.RESULTS.clicked.connect(self.ResultFrame)
        self.HOME.clicked.connect(self.HomeFrame)
        self.LOGOUT.clicked.connect(self.Logout)
        self.TAKEEXAM.clicked.connect(self.Exam)


    def ExamFrame(self):
        self.EXAMframe.show()
        self.HOMEframe.hide()
        self.StudentResults.hide()

    def ResultFrame(self):
        self.EXAMframe.hide()
        self.HOMEframe.hide()
        self.StudentResults.show()


    def HomeFrame(self):
        self.EXAMframe.hide()
        self.HOMEframe.show()
        self.StudentResults.hide()

    def Exam(self):
        prompt = QMessageBox.question(self, "NOTICE",
                                      "DO YOU WANT TO TAKE THE EXAM NOW? 'ARE YOU SURE?",
                                      QMessageBox.Yes | QMessageBox.No)
        if prompt == QMessageBox.Yes:
            Doc = Document("Exam_file.docx")
            Doc.save("Exam_file(NEW).docx")
            os.startfile("Exam_file.docx")

        else:
            QMessageBox.about(self,"NOTICE","IF YOU DO NOT WANT TO TAKE YOUR EXAM FOR NOW PLEASE SET A DATE FOR YOUR EXAMINATION, THANK YOU FOR UNDERSTANDING")

    def Logout(self):
        log.show()
        estudyante.hide()

class Main(Main, QMainWindow):
    def __init__(self):
        super(Main, self).__init__()
        self.setupUi(self)

        self.AdminResults.hide()
        self.AdminAddExam.hide()
        self.BackupFrame.hide()
        self.ReportFrame.hide()
        self.HomeAdmin.show()
        current = datetime.datetime.now()
        self.DATE_2.setText(str(current))

        self.HOME.clicked.connect(self.HOMEadmin)
        self.EXAM.clicked.connect(self.ExamAdmin)
        self.RESULTS.clicked.connect(self.ResultsAdmin)
        self.BACKUP.clicked.connect(self.BackupAdmin)
        self.REPORT.clicked.connect(self.ReportAdmin)
        self.SEARCH.clicked.connect(self.Search)
        #self.SaveData_2.clicked.connect(self.Delete)
        self.SaveData.clicked.connect(self.Excel)
        self.ADDEXAM.clicked.connect(self.Open)
        self.LOGOUT.clicked.connect(self.Logout)
        self.CONFIRM.clicked.connect(self.Setscore)
        self.HIGHEST.clicked.connect(self.SortHighest)
        self.LOWEST.clicked.connect(self.SortLowest)
        self.PASS.clicked.connect(self.SortPass)
        self.FAIL.clicked.connect(self.SortFail)
        self.VIEWEXAMINEES.clicked.connect(self.ViewExaminees)
        self.SHOWEXCEL.clicked.connect(self.openExcel)

        self.cone = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")



    def HOMEadmin(self):
        self.HomeAdmin.show()
        self.AdminAddExam.hide()
        self.AdminResults.hide()
        self.BackupFrame.hide()
        self.ReportFrame.hide()

    def ExamAdmin(self):
        self.HomeAdmin.hide()
        self.AdminAddExam.show()
        self.AdminResults.hide()
        self.BackupFrame.hide()
        self.ReportFrame.hide()

    def ResultsAdmin(self):
        conn = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")
        cursor = conn.cursor()
        cursor.execute("select lastname,firstname,middlename,shift,username,passwd,stat,score from naragdata;")
        self.ResultsTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.ResultsTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.ResultsTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))

        conn.close()
        self.HomeAdmin.hide()
        self.AdminAddExam.hide()
        self.AdminResults.show()
        self.BackupFrame.hide()
        self.ReportFrame.hide()

    def ViewExaminees(self):
        exam.show()
    def BackupAdmin(self):
        conn = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")
        cursor = conn.cursor()
        cursor.execute("select lastname,firstname,middlename,shift,username,passwd,stat,score from naragdata;")
        self.BackupTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.BackupTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.BackupTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))

        self.HomeAdmin.hide()
        self.AdminAddExam.hide()
        self.AdminResults.hide()
        self.BackupFrame.show()
        self.ReportFrame.hide()

    def ReportAdmin(self):
        self.HomeAdmin.hide()
        self.AdminAddExam.hide()
        self.AdminResults.hide()
        self.BackupFrame.hide()
        self.ReportFrame.show()

    def Search(self):
        line = self.lineEdit.text()
        cursor1 = self.cone.cursor()

        query = "select * from naragdata where lastname like '%"+line+"%' ;"
        data = (line)
        cursor1.execute(query,data)

        self.ResultsTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor1):
            self.ResultsTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.ResultsTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))


    def SortHighest(self):
        cursor = self.cone.cursor()
        query = "select * from naragdata order by score desc;"
        cursor.execute(query)

        self.ResultsTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.ResultsTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.ResultsTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))

    def SortLowest(self):
        cursor = self.cone.cursor()
        query = "select * from naragdata order by score asc;"
        cursor.execute(query)

        self.ResultsTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.ResultsTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.ResultsTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))

    def SortPass(self):
        cursor = self.cone.cursor()
        query = "select * from naragdata where stat='Passed';"
        cursor.execute(query)

        self.ResultsTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.ResultsTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.ResultsTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))

    def SortFail(self):
        cursor = self.cone.cursor()
        query = "select * from naragdata where stat='Failed' or 'None';"
        cursor.execute(query)

        self.ResultsTable.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.ResultsTable.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.ResultsTable.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))

    def Setscore(self):
        global status
        Name = self.NAMES.text()
        First = self.FIRSTNAME.text()
        Score = self.SCORE.text()
        cursor = self.cone.cursor()
        if self.PASS_2.isChecked()==True:
            status1="Passed"
        elif self.FAIL_2.isChecked()==True:
            status1="Failed"


        query = "select * from naragdata where lastname=%s and firstname=%s;"
        data = (Name,First)
        cursor.execute(query,data)
        a = cursor.fetchall()
        meet = len(a)
        self.cone.commit()
        if meet > 0:
            query1 = "update naragdata set score=%s,stat=%s where lastname=%s and firstname=%s;"
            data1 = (Score, status1, Name, First)
            cursor.execute(query1, data1)
            QMessageBox.about(self, 'Success', 'Successfully Added')
            self.NAMES.clear()
            self.FIRSTNAME.clear()
            self.SCORE.clear()
            self.cone.commit()

        elif Name == '' or First == '' or Score == '':
            QMessageBox.warning(self, 'Error', 'Please fill up all the necessary informations')
        else:
            QMessageBox.warning(self, 'Name could not find', 'Check the database!')


    def Open(self):
        prompt = QMessageBox.question(self,"NOTICE","Do you want to create a new examination? 'click yes' and if you want to view the latest exam 'click no'",QMessageBox.Yes | QMessageBox.No)
        if prompt == QMessageBox.Yes:
            Doc = Document()
            Doc.save("Exam_file.docx")
            os.startfile("Exam_file.docx")
        else:
            Doc = Document("Exam_file.docx")
            Doc.save("Exam_file.docx")
            os.startfile("Exam_file.docx")

    def Excel(self):
        cur = self.cone.cursor()
        query = "select * from naragdata;"
        cur.execute(query)
        gather = cur.fetchall()

        workbook = xlsxwriter.Workbook('Data.xlsx')
        worksheet = workbook.add_worksheet()

        for i,row in enumerate(gather):
            for c,col in enumerate(row):
                worksheet.write(i,c,col)

        QMessageBox.about(self, "YOUR DATA HAS BEEN SAVED",
                          "SUCCESSFULLY SAVED IN EXCEL!")

        workbook.close()

    def openExcel(self):
        os.startfile("Data.xlsx")


    def Logout(self):
        adm.show()
        system.hide()


class Exam(Exam, QMainWindow):
    def __init__(self):
        super(Exam, self).__init__()

        self.setupUi(self)
        conn = mysql.connector.connect(host="localhost", user="root", passwd="narag", database="preboard")
        cursor = conn.cursor()
        cursor.execute("select lastname,firstname,middlename from naragdata;")
        self.tableWidget.setRowCount(0)

        for row_number, row_data in enumerate(cursor):
            self.tableWidget.insertRow(row_number)
            for colum_number, data in enumerate(row_data):
                self.tableWidget.setItem(row_number, colum_number, QtWidgets.QTableWidgetItem(str(data)))



if __name__ == "__main__":
    F = QApplication(sys.argv)
    una = First()
    una.show()
    adm = Admin()
    adm.hide()
    log = Login()
    log.hide()
    sign = Signup()
    sign.hide()
    estudyante = Student()
    estudyante.hide()
    system = Main()
    system.hide()
    exam = Exam()
    exam.hide()

    sys.exit(F.exec())
